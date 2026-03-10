import streamlit as st
import time
from docx import Document
import io
import streamlit.components.v1 as components

st.set_page_config(page_title="Ficha Cívica - Salud Pública", page_icon="⚖️", layout="centered")

# --- SISTEMA DE TIEMPO INTELIGENTE Y BOTÓN GO ---
if 'ficha_iniciada_civica' not in st.session_state:
    st.session_state.ficha_iniciada_civica = False

if not st.session_state.ficha_iniciada_civica:
    st.info("👋 ¡Hola! Tienes 20 minutos para resolver esta ficha. El tiempo comenzará a correr cuando presiones el botón de abajo.")
    col_btn1, col_btn2, col_btn3 = st.columns([1, 2, 1])
    with col_btn2:
        if st.button("🚀 ESTOY LISTO: INICIAR FICHA", use_container_width=True):
            st.session_state.ficha_iniciada_civica = True
            st.session_state.inicio_tiempo_civica = time.time()
            st.session_state.minutos_asignados_civica = 20
            st.rerun()
    st.stop()

segundos_transcurridos = time.time() - st.session_state.inicio_tiempo_civica
segundos_restantes = (st.session_state.minutos_asignados_civica * 60) - segundos_transcurridos
tiempo_agotado = segundos_restantes <= 0
bloquear_inputs = tiempo_agotado 

# --- MENÚ LATERAL: RELOJ VISUAL EN TIEMPO REAL ---
with st.sidebar:
    st.markdown("### ⏱️ Cronómetro de ficha")
    if not tiempo_agotado:
        reloj_html = f"""
        <div style="background-color: #f0f2f6; padding: 15px; border-radius: 10px; text-align: center; border: 2px solid #e0e4eb;">
            <h2 id="reloj" style="margin: 0; color: #2ecc71; font-family: monospace; font-size: 38px;">--:--</h2>
        </div>
        <script>
            var tiempo = {int(segundos_restantes)};
            var display = document.getElementById('reloj');
            var intervalo = setInterval(function() {{
                if (tiempo <= 0) {{
                    clearInterval(intervalo);
                    display.innerHTML = "00:00";
                    display.style.color = "#e74c3c";
                }} else {{
                    var min = Math.floor(tiempo / 60).toString().padStart(2, '0');
                    var sec = (tiempo % 60).toString().padStart(2, '0');
                    display.innerHTML = min + ":" + sec;
                    if (tiempo <= 300) display.style.color = "#f39c12"; // Amarillo últimos 5 min
                    if (tiempo <= 60) display.style.color = "#e74c3c"; // Rojo último minuto
                    tiempo--;
                }}
            }}, 1000);
        </script>
        """
        components.html(reloj_html, height=85)
        
        if st.button("➕ Dar 4 min extra", use_container_width=True):
            st.session_state.minutos_asignados_civica += 4
            st.rerun()
        st.caption("El reloj avanza en tiempo real. Al llegar a cero, haz clic en cualquier parte para bloquear y descargar tu avance.")
    else:
        st.error("## 00:00\n⚠️ TIEMPO AGOTADO")
        st.write("Tu ficha ha sido bloqueada. Por favor, descarga tu avance en la parte inferior.")
        
        if st.button("🔓 Desbloquear (Dar 4 min extra)", use_container_width=True):
            st.session_state.minutos_asignados_civica += 4
            st.rerun()
# --------------------------------------

st.markdown("### FICHA DE CÍVICA – SESIÓN 2° AÑO A - B - C")
st.markdown("**Tema:** Derecho a la Vida y la Salud Pública")
st.markdown("**Prof:** Eduardo Florez Montero / Unidad 1")
st.markdown("---")

col1, col2, col3 = st.columns([3, 1, 2])
with col1: nombre = st.text_input("Estudiante:", disabled=bloquear_inputs)
with col2: seccion = st.selectbox("Sección:", ["", "A", "B", "C"], disabled=bloquear_inputs)
with col3: fecha = st.text_input("Fecha:", disabled=bloquear_inputs)

st.markdown("---")
with st.expander("🎯 META DE HOY", expanded=True): st.write("Relacionar el derecho a la vida con la salud pública y proponer normas de prevención e información responsable.")

st.markdown("---")
st.subheader("📚 Material de Apoyo")
try:
    with open("Presentacion_Civica.pdf", "rb") as file:
        st.download_button("📄 Descargar Presentación de la Clase (PDF)", data=file, file_name="Presentacion_Civica.pdf", mime="application/pdf")
except FileNotFoundError:
    st.info("📌 La presentación de la clase se está procesando y aparecerá aquí en breve.")

st.markdown("---")
st.subheader("NIVEL 1 (FÁCIL) – PARA TODOS")
st.write("Empecemos identificando acciones cotidianas que protegen nuestra comunidad.")
col_n1a, col_n1b = st.columns(2)
with col_n1a: q1_1 = st.text_input("1.1) Primera acción concreta de prevención:", disabled=bloquear_inputs)
with col_n1b: q1_2 = st.text_input("1.2) Segunda acción de prevención:", disabled=bloquear_inputs)

q2 = st.text_input("2) Completa: Si una persona difunde información falsa sobre temas de salud, esto puede ocasionar que...", disabled=bloquear_inputs)
q3 = st.radio("3) Antes de compartir una noticia sobre salud en mis redes sociales, lo primero que debo hacer es:", ["verificar si la fuente es confiable", "reenviar rápidamente sin leer", "insultar a quien lo publicó"], horizontal=True, disabled=bloquear_inputs)
q4 = st.text_input("4) Escribe un compromiso personal que puedas aplicar dentro del aula para mantener una buena salud grupal:", disabled=bloquear_inputs)

st.markdown("---")
st.subheader("NIVEL 2 (MEDIO) – RETO 1")
st.write("5) Lee el siguiente caso: Ves una publicación en redes que dice 'Las vacunas hacen daño' pero no muestra ninguna fuente científica. Escribe dos preguntas críticas que te harías para verificar si esa información es falsa:")
q5_1 = st.text_input("Pregunta crítica 1:", disabled=bloquear_inputs)
q5_2 = st.text_input("Pregunta crítica 2:", disabled=bloquear_inputs)

st.write("6) Escribe dos reglas de oro que todos deberíamos seguir antes de compartir información de salud:")
q6_1 = st.text_input("Regla de oro 1:", disabled=bloquear_inputs)
q6_2 = st.text_input("Regla de oro 2:", disabled=bloquear_inputs)

st.markdown("---")
st.subheader("NIVEL 3 (DIFÍCIL) – RETO 2")
st.write("7) Imagina que eres el líder de una campaña de salud escolar. Crea un mensaje corto y llamativo (máximo 2 líneas) para animar a tus compañeros a prevenir enfermedades:")
q7 = st.text_area("Mensaje de campaña:", disabled=bloquear_inputs)

st.write("8) Redacta un párrafo corto (aprox. 5 líneas) explicando con tus propias palabras: ¿Por qué crees que informarnos responsablemente es una forma de proteger nuestra vida y la de los demás?")
q8 = st.text_area("Tu reflexión:", disabled=bloquear_inputs)

st.markdown("---")
st.subheader("📊 Valoración de la Actividad")
val_funcional = st.slider("1. ¿Qué tan fácil y funcional te pareció usar esta ficha digital?", 1, 5, 5, disabled=bloquear_inputs)
val_interes = st.radio("2. ¿El tema y las actividades te parecieron interesantes?", ["Sí, mucho", "Estuvo bien", "No mucho", "Nada interesante"], horizontal=True, disabled=bloquear_inputs)

st.markdown("---")
st.subheader("📋 Lista de Cotejo (Uso exclusivo del docente)")
st.caption("Estos son los criterios con los que tu profesor evaluará esta ficha:")
st.checkbox("Identifica y describe correctamente dos acciones prácticas de prevención (Nivel 1).", value=False, disabled=True)
st.checkbox("Analiza críticamente información de salud formulando preguntas pertinentes (Nivel 2).", value=False, disabled=True)
st.checkbox("Propone normas claras y responsables para compartir información en redes (Nivel 2).", value=False, disabled=True)
st.checkbox("Redacta un mensaje de campaña claro y argumenta sólidamente su reflexión final (Nivel 3).", value=False, disabled=True)

st.markdown("---")

if st.button("Generar mi Evidencia en Word"):
    if not nombre.strip() or seccion == "":
        st.error("⚠️ Por favor, ingresa tu nombre y selecciona tu sección para poder identificarte.")
    elif not tiempo_agotado and (
        not q1_1.strip() or not q1_2.strip() or not q2.strip() or not q4.strip() or 
        not q5_1.strip() or not q5_2.strip() or not q6_1.strip() or not q6_2.strip() or
        not q7.strip() or not q8.strip()
    ):
        st.error("⚠️ Aún tienes tiempo. Debes completar las preguntas de **TODOS LOS NIVELES (1, 2 y 3)** antes de descargar tu evidencia.")
    else:
        doc = Document()
        doc.add_heading('FICHA DE CÍVICA – SESIÓN 2° AÑO', level=1)
        doc.add_paragraph(f'Estudiante: {nombre} | Sección: {seccion} | Fecha: {fecha}')
        if tiempo_agotado:
            doc.add_paragraph('[Entregado al finalizar el tiempo reglamentario]').bold = True
        doc.add_paragraph('---')
        
        doc.add_heading('NIVEL 1 (FÁCIL)', level=2)
        p1 = doc.add_paragraph()
        p1.add_run('1) Acciones de prevención:\n').bold = True
        p1.add_run(f'• {q1_1}\n• {q1_2}\n\n')
        p1.add_run('2) Difundir información falsa puede:\n').bold = True
        p1.add_run(f'• {q2}\n\n')
        p1.add_run('3) Antes de compartir una noticia debo:\n').bold = True
        p1.add_run(f'• {q3}\n\n')
        p1.add_run('4) Compromiso de salud en el aula:\n').bold = True
        p1.add_run(f'• {q4}')
        
        doc.add_heading('NIVEL 2 (MEDIO)', level=2)
        p2 = doc.add_paragraph()
        p2.add_run('5) Preguntas críticas (Caso Vacunas):\n').bold = True
        p2.add_run(f'• {q5_1}\n• {q5_2}\n\n')
        p2.add_run('6) Reglas de oro para compartir información:\n').bold = True
        p2.add_run(f'• {q6_1}\n• {q6_2}')
        
        doc.add_heading('NIVEL 3 (DIFÍCIL)', level=2)
        p3 = doc.add_paragraph()
        p3.add_run('7) Mensaje de campaña:\n').bold = True
        p3.add_run(f'• {q7}\n\n')
        p3.add_run('8) Reflexión sobre la información responsable:\n').bold = True
        p3.add_run(f'• {q8}')
        
        doc.add_heading('Valoración de la Actividad', level=2)
        p4 = doc.add_paragraph()
        p4.add_run('Funcionalidad de la ficha digital: ').bold = True
        p4.add_run(f'{val_funcional} estrellas\n')
        p4.add_run('Interés en el tema: ').bold = True
        p4.add_run(f'{val_interes}')
        
        doc.add_page_break()
        doc.add_heading('Lista de Cotejo - Evaluación del Docente', level=2)
        doc.add_paragraph('[ ] Identifica y describe correctamente dos acciones prácticas de prevención (Nivel 1).')
        doc.add_paragraph('[ ] Analiza críticamente información de salud formulando preguntas pertinentes (Nivel 2).')
        doc.add_paragraph('[ ] Propone normas claras y responsables para compartir información en redes (Nivel 2).')
        doc.add_paragraph('[ ] Redacta un mensaje de campaña claro y argumenta sólidamente su reflexión final (Nivel 3).')
        doc.add_paragraph('\nNota / Observaciones: ________________________________________________')

        bio = io.BytesIO()
        doc.save(bio)
        
        if not tiempo_agotado: st.balloons()
        st.success("¡Tu archivo está listo para entregar!")
        st.download_button(label="📥 Descargar Documento Final (.docx)", data=bio.getvalue(), file_name=f"Ficha_Civica_2{seccion}_{nombre.replace(' ', '_')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
