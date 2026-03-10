import streamlit as st
import time
from docx import Document
import io
import streamlit.components.v1 as components

st.set_page_config(page_title="Ficha PFRH 1ro - Cambios", page_icon="🌱", layout="centered")

# --- SISTEMA DE TIEMPO INTELIGENTE Y BOTÓN GO ---
if 'ficha_iniciada_pfrh1' not in st.session_state:
    st.session_state.ficha_iniciada_pfrh1 = False

if not st.session_state.ficha_iniciada_pfrh1:
    st.info("👋 ¡Hola! Tienes 20 minutos para resolver esta ficha. El tiempo comenzará a correr cuando presiones el botón de abajo.")
    col_btn1, col_btn2, col_btn3 = st.columns([1, 2, 1])
    with col_btn2:
        if st.button("🚀 ESTOY LISTO: INICIAR FICHA", use_container_width=True):
            st.session_state.ficha_iniciada_pfrh1 = True
            st.session_state.inicio_tiempo_pfrh1 = time.time()
            st.session_state.minutos_asignados_pfrh1 = 20
            st.rerun()
    st.stop()

segundos_transcurridos = time.time() - st.session_state.inicio_tiempo_pfrh1
segundos_restantes = (st.session_state.minutos_asignados_pfrh1 * 60) - segundos_transcurridos
tiempo_agotado = segundos_restantes <= 0
bloquear_inputs = tiempo_agotado 

# --- MENÚ LATERAL: RELOJ VISUAL EN TIEMPO REAL ---
with st.sidebar:
    st.markdown("### ⏱️ Cronómetro de ficha")
    if not tiempo_agotado:
        reloj_html = f"""
        <div style="background-color: #f0f2f6; padding: 15px; border-radius: 10px; text-align: center; border: 2px solid #e0e4eb;">
            <h2 id="reloj" style="margin: 0; color: #2ecc71; font-family: monospace; font-size: 38px;">--:--</h2>
        </div>
        <script>
            var tiempo = {int(segundos_restantes)};
            var display = document.getElementById('reloj');
            var intervalo = setInterval(function() {{
                if (tiempo <= 0) {{
                    clearInterval(intervalo);
                    display.innerHTML = "00:00";
                    display.style.color = "#e74c3c";
                }} else {{
                    var min = Math.floor(tiempo / 60).toString().padStart(2, '0');
                    var sec = (tiempo % 60).toString().padStart(2, '0');
                    display.innerHTML = min + ":" + sec;
                    if (tiempo <= 300) display.style.color = "#f39c12"; // Amarillo últimos 5 min
                    if (tiempo <= 60) display.style.color = "#e74c3c"; // Rojo último minuto
                    tiempo--;
                }}
            }}, 1000);
        </script>
        """
        components.html(reloj_html, height=85)
        
        if st.button("➕ Dar 4 min extra", use_container_width=True):
            st.session_state.minutos_asignados_pfrh1 += 4
            st.rerun()
        st.caption("El reloj avanza en tiempo real. Al llegar a cero, haz clic en cualquier parte para bloquear y descargar tu avance.")
    else:
        st.error("## 00:00\n⚠️ TIEMPO AGOTADO")
        st.write("Tu ficha ha sido bloqueada. Por favor, descarga tu avance en la parte inferior.")
        
        if st.button("🔓 Desbloquear (Dar 4 min extra)", use_container_width=True):
            st.session_state.minutos_asignados_pfrh1 += 4
            st.rerun()
# --------------------------------------

st.markdown("### FICHA DE PFRH – SESIÓN 1° AÑO A - B - C")
st.markdown("**Tema:** Cambios físicos y emocionales en la adolescencia")
st.markdown("**Prof:** Eduardo Florez Montero / Unidad 1")
st.markdown("---")

col1, col2, col3 = st.columns([3, 1, 2])
with col1: nombre = st.text_input("Estudiante:", disabled=bloquear_inputs)
with col2: seccion = st.selectbox("Sección:", ["", "A", "B", "C"], disabled=bloquear_inputs)
with col3: fecha = st.text_input("Fecha:", disabled=bloquear_inputs)

st.warning("🔒 **Seguridad:** Recuerda proteger tu privacidad. Al responder, usa ejemplos generales sin dar nombres reales de compañeros ni datos personales.")

st.markdown("---")
with st.expander("🎯 META DE HOY", expanded=True): st.write("Identificar cambios físicos y emocionales de la adolescencia y aplicar estrategias de autorregulación.")

st.markdown("---")
st.subheader("📚 Material de Apoyo")
try:
    with open("Presentacion_PFRH_1ro.pdf", "rb") as file:
        st.download_button("📄 Descargar Presentación de la Clase (PDF)", data=file, file_name="Presentacion_PFRH_1ro.pdf", mime="application/pdf")
except FileNotFoundError:
    st.info("📌 La presentación de la clase se está procesando y aparecerá aquí en breve.")

st.markdown("---")
st.subheader("NIVEL 1 (FÁCIL) – PARA TODOS")
st.write("¡Empecemos observando cómo crecemos! Escribe dos cambios que hayas notado o aprendido sobre la adolescencia:")
col_f1, col_f2 = st.columns(2)
with col_f1: 
    q1_1 = st.text_input("1.1) Un cambio físico:", disabled=bloquear_inputs)
    q2_1 = st.text_input("2.1) Un cambio emocional:", disabled=bloquear_inputs)
with col_f2: 
    q1_2 = st.text_input("1.2) Otro cambio físico:", disabled=bloquear_inputs)
    q2_2 = st.text_input("2.2) Otro cambio emocional:", disabled=bloquear_inputs)

q3 = st.text_input("3) Completa la siguiente frase: 'Cuando siento mucho enojo y siento que voy a explotar, yo puedo...'", disabled=bloquear_inputs)
q4 = st.radio("4) Frente a un momento de estrés, ¿cuál de estas es una reacción saludable?", ["Gritar sin control", "Respirar profundo y hablar", "Insultar en mis redes sociales"], horizontal=True, disabled=bloquear_inputs)

st.markdown("---")
st.subheader("NIVEL 2 (MEDIO) – RETO 1")
st.write("5) Analiza este caso: Estás en el recreo y alguien hace un comentario desagradable sobre tu cuerpo o apariencia. Identifica los elementos de esta situación:")
col_c1, col_c2, col_c3 = st.columns(3)
with col_c1: q5_emo = st.text_input("¿Qué emoción sentirías?:", disabled=bloquear_inputs)
with col_c2: q5_cau = st.text_input("¿Cuál fue la causa directa?:", disabled=bloquear_inputs)
with col_c3: q5_con = st.text_input("¿Cuál podría ser la consecuencia?:", disabled=bloquear_inputs)

st.markdown("---")
st.subheader("NIVEL 3 (DIFÍCIL) – RETO 2")
st.write("6) Construye tu propio **Plan de Emergencia Emocional**. Escribe 3 pasos que harías para calmarte antes de responderle a alguien que te hizo enojar mucho:")
q6_1 = st.text_input("Paso 1:", disabled=bloquear_inputs)
q6_2 = st.text_input("Paso 2:", disabled=bloquear_inputs)
q6_3 = st.text_input("Paso 3:", disabled=bloquear_inputs)

st.write("7) Imagina que lograste calmarte. Escribe un mensaje asertivo (claro, directo, pero sin usar ningún insulto) que le dirías a esa persona para poner un límite:")
q7 = st.text_area("Tu mensaje asertivo:", disabled=bloquear_inputs)

st.markdown("---")
# --- VALORACIÓN DEL ESTUDIANTE ---
st.subheader("📊 Valoración de la Actividad")
val_funcional = st.slider("1. ¿Qué tan fácil y funcional te pareció usar esta ficha digital?", 1, 5, 5, disabled=bloquear_inputs)
val_interes = st.radio("2. ¿El tema y las actividades te parecieron interesantes?", ["Sí, mucho", "Estuvo bien", "No mucho", "Nada interesante"], horizontal=True, disabled=bloquear_inputs)

st.markdown("---")
# --- LISTA DE COTEJO PARA EL DOCENTE ---
st.subheader("📋 Lista de Cotejo (Uso exclusivo del docente)")
st.caption("Estos son los criterios con los que tu profesor evaluará esta ficha:")
st.checkbox("Identifica correctamente cambios físicos y emocionales propios de la adolescencia (Nivel 1).", value=False, disabled=True)
st.checkbox("Reconoce estrategias básicas de regulación emocional y reacciones saludables (Nivel 1).", value=False, disabled=True)
st.checkbox("Analiza adecuadamente la emoción, causa y consecuencia ante un caso de estrés o burla (Nivel 2).", value=False, disabled=True)
st.checkbox("Propone un plan de autorregulación de 3 pasos y redacta un mensaje asertivo sin agresiones (Nivel 3).", value=False, disabled=True)

st.markdown("---")

if st.button("Generar mi Evidencia en Word"):
    if not nombre.strip() or seccion == "":
        st.error("⚠️ Por favor, ingresa tu nombre y selecciona tu sección para poder identificarte.")
        
    elif not tiempo_agotado and (
        not q1_1.strip() or not q1_2.strip() or not q2_1.strip() or not q2_2.strip() or not q3.strip() or
        not q5_emo.strip() or not q5_cau.strip() or not q5_con.strip() or
        not q6_1.strip() or not q6_2.strip() or not q6_3.strip() or not q7.strip()
    ):
        st.error("⚠️ Aún tienes tiempo. Debes completar las preguntas de **TODOS LOS NIVELES (1, 2 y 3)** antes de descargar tu evidencia.")
        
    else:
        doc = Document()
        
        doc.add_heading('FICHA DE PFRH – SESIÓN 1° AÑO', level=1)
        doc.add_paragraph(f'Estudiante: {nombre} | Sección: {seccion} | Fecha: {fecha}')
        if tiempo_agotado:
            doc.add_paragraph('[Entregado al finalizar el tiempo reglamentario]').bold = True
        doc.add_paragraph('---')
        
        doc.add_heading('NIVEL 1 (FÁCIL)', level=2)
        p1 = doc.add_paragraph()
        p1.add_run('1) Cambios físicos:\n').bold = True
        p1.add_run(f'• {q1_1}\n• {q1_2}\n\n')
        p1.add_run('2) Cambios emocionales:\n').bold = True
        p1.add_run(f'• {q2_1}\n• {q2_2}\n\n')
        p1.add_run('3) Cuando siento enojo, puedo:\n').bold = True
        p1.add_run(f'• {q3}\n\n')
        p1.add_run('4) Reacción saludable:\n').bold = True
        p1.add_run(f'• {q4}')
        
        doc.add_heading('NIVEL 2 (MEDIO)', level=2)
        p2 = doc.add_paragraph()
        p2.add_run('5) Caso "Comentario sobre mi cuerpo":\n').bold = True
        p2.add_run(f'• Emoción: {q5_emo}\n')
        p2.add_run(f'• Causa: {q5_cau}\n')
        p2.add_run(f'• Consecuencia: {q5_con}')
        
        doc.add_heading('NIVEL 3 (DIFÍCIL)', level=2)
        p3 = doc.add_paragraph()
        p3.add_run('6) Plan de autorregulación en 3 pasos:\n').bold = True
        p3.add_run(f'• Paso 1: {q6_1}\n• Paso 2: {q6_2}\n• Paso 3: {q6_3}\n\n')
        p3.add_run('7) Mensaje asertivo (Límites):\n').bold = True
        p3.add_run(f'• {q7}')
        
        doc.add_heading('Valoración de la Actividad', level=2)
        p4 = doc.add_paragraph()
        p4.add_run('Funcionalidad de la ficha digital: ').bold = True
        p4.add_run(f'{val_funcional} estrellas\n')
        p4.add_run('Interés en el tema: ').bold = True
        p4.add_run(f'{val_interes}')
        
        doc.add_page_break()
        
        doc.add_heading('Lista de Cotejo - Evaluación del Docente', level=2)
        doc.add_paragraph('[ ] Identifica correctamente cambios físicos y emocionales propios de la adolescencia (Nivel 1).')
        doc.add_paragraph('[ ] Reconoce estrategias básicas de regulación emocional y reacciones saludables (Nivel 1).')
        doc.add_paragraph('[ ] Analiza adecuadamente la emoción, causa y consecuencia ante un caso de estrés o burla (Nivel 2).')
        doc.add_paragraph('[ ] Propone un plan de autorregulación de 3 pasos y redacta un mensaje asertivo sin agresiones (Nivel 3).')
        doc.add_paragraph('\nNota / Observaciones: ________________________________________________')
        
        bio = io.BytesIO()
        doc.save(bio)
        
        if not tiempo_agotado: st.balloons()
            
        st.success("¡Tu archivo está listo para entregar!")
        st.download_button(
            label="📥 Descargar Documento Final (.docx)", 
            data=bio.getvalue(), 
            file_name=f"Ficha_PFRH_1{seccion}_{nombre.replace(' ', '_')}.docx", 
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
