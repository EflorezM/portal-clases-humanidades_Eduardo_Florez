import streamlit as st
import time
from docx import Document
import io
import streamlit.components.v1 as components

st.set_page_config(page_title="Ficha de Religión - El judaísmo", page_icon="📜", layout="centered")

# --- SISTEMA DE TIEMPO INTELIGENTE Y BOTÓN GO ---
if 'ficha_iniciada_religion' not in st.session_state:
    st.session_state.ficha_iniciada_religion = False

if not st.session_state.ficha_iniciada_religion:
    st.info("👋 ¡Hola! Tienes 20 minutos para resolver esta ficha. El tiempo comenzará a correr cuando presiones el botón de abajo.")
    col_btn1, col_btn2, col_btn3 = st.columns([1, 2, 1])
    with col_btn2:
        if st.button("🚀 ESTOY LISTO: INICIAR FICHA", use_container_width=True):
            st.session_state.ficha_iniciada_religion = True
            st.session_state.inicio_tiempo_religion = time.time()
            st.session_state.minutos_asignados_religion = 20
            st.rerun()
    st.stop()

segundos_transcurridos = time.time() - st.session_state.inicio_tiempo_religion
segundos_restantes = (st.session_state.minutos_asignados_religion * 60) - segundos_transcurridos
tiempo_agotado = segundos_restantes <= 0
bloquear_inputs = tiempo_agotado 

# --- MENÚ LATERAL: RELOJ VISUAL EN TIEMPO REAL ---
with st.sidebar:
    st.markdown("### ⏱️ Cronómetro de ficha")
    if not tiempo_agotado:
        reloj_html = f"""
        <div style="background-color: #f0f2f6; padding: 15px; border-radius: 10px; text-align: center; border: 2px solid #e0e4eb;">
            <h2 id="reloj" style="margin: 0; color: #2ecc71; font-family: monospace; font-size: 38px;">--:--</h2>
        </div>
        <script>
            var tiempo = {int(segundos_restantes)};
            var display = document.getElementById('reloj');
            var intervalo = setInterval(function() {{
                if (tiempo <= 0) {{
                    clearInterval(intervalo);
                    display.innerHTML = "00:00";
                    display.style.color = "#e74c3c";
                }} else {{
                    var min = Math.floor(tiempo / 60).toString().padStart(2, '0');
                    var sec = (tiempo % 60).toString().padStart(2, '0');
                    display.innerHTML = min + ":" + sec;
                    if (tiempo <= 300) display.style.color = "#f39c12"; // Amarillo
                    if (tiempo <= 60) display.style.color = "#e74c3c"; // Rojo
                    tiempo--;
                }}
            }}, 1000);
        </script>
        """
        components.html(reloj_html, height=85)
        
        if st.button("➕ Dar 4 min extra", use_container_width=True):
            st.session_state.minutos_asignados_religion += 4
            st.rerun()
        st.caption("El reloj avanza en tiempo real. Al llegar a cero, haz clic para bloquear y descargar tu avance.")
    else:
        st.error("## 00:00\n⚠️ TIEMPO AGOTADO")
        st.write("Tu ficha ha sido bloqueada. Por favor, descarga tu avance en la parte inferior.")
        
        if st.button("🔓 Desbloquear (Dar 4 min extra)", use_container_width=True):
            st.session_state.minutos_asignados_religion += 4
            st.rerun()
# --------------------------------------

st.markdown("### FICHA DE RELIGIÓN – SESIÓN 1° AÑO A - B - C")
st.markdown("**Tema:** El judaísmo: Cuna de las religiones monoteístas.")
st.markdown("**Prof:** Eduardo Florez Montero / Unidad 1")
st.markdown("---")

col1, col2, col3 = st.columns([3, 1, 2])
with col1: nombre = st.text_input("Estudiante:", disabled=bloquear_inputs)
with col2: seccion = st.selectbox("Sección:", ["", "A", "B", "C"], disabled=bloquear_inputs)
with col3: fecha = st.text_input("Fecha:", disabled=bloquear_inputs)

st.markdown("---")
with st.expander("🎯 META DE HOY", expanded=True): st.write("Reconocer aportes del judaísmo y practicar diálogo respetuoso sin prejuicios.")

st.markdown("---")
st.subheader("📚 Material de Apoyo")
try:
    with open("Presentacion_Religion.pdf", "rb") as file:
        st.download_button("📄 Descargar Presentación de la Clase (PDF)", data=file, file_name="Presentacion_Religion.pdf", mime="application/pdf")
except FileNotFoundError:
    st.info("📌 La presentación de la clase se está procesando y aparecerá aquí en breve.")

st.markdown("---")
st.subheader("NIVEL 1 (FÁCIL) – PARA TODOS")
st.write("¡Comencemos recordando los conceptos clave de la clase de hoy!")
q1 = st.text_input("1) Piensa en la palabra 'monoteísmo'. Según lo que vimos, ¿qué significa exactamente creer en esto?", disabled=bloquear_inputs)
q2 = st.radio("2) Selecciona la opción correcta: La figura histórica de Abraham es importante para...", ["judaísmo", "cristianismo", "ambos"], horizontal=True, disabled=bloquear_inputs)

st.write("3) Abraham demostró muchas cualidades. Menciona 2 valores importantes que podemos aprender de su historia:")
col_v1, col_v2 = st.columns(2)
with col_v1: q3_1 = st.text_input("Valor 1:", disabled=bloquear_inputs)
with col_v2: q3_2 = st.text_input("Valor 2:", disabled=bloquear_inputs)

q4 = st.text_input("4) Escribe una regla de oro que debes usar siempre para dialogar con respeto sobre las creencias de otras personas:", disabled=bloquear_inputs)

st.markdown("---")
st.subheader("NIVEL 2 (MEDIO) – RETO 1")
st.info("💡 **Reto 1:** Piensa en lo que hemos aprendido en clase. ¿En qué se parecen el judaísmo y el cristianismo? Escribe dos características que ambas religiones compartan (semejanzas):")
col_s1, col_s2 = st.columns(2)
with col_s1: q5_1 = st.text_input("Primera semejanza:", disabled=bloquear_inputs)
with col_s2: q5_2 = st.text_input("Segunda semejanza:", disabled=bloquear_inputs)

q6 = st.text_input("6) Para mantener una buena convivencia, escribe un prejuicio o idea equivocada sobre otras religiones que debemos evitar en el salón de clases:", disabled=bloquear_inputs)

st.markdown("---")
st.subheader("NIVEL 3 (DIFÍCIL) – RETO 2")
q7 = st.text_area("7) Reflexiona a profundidad (aprox. 5 líneas): ¿Por qué crees que el respeto hacia otras creencias es una parte fundamental de los valores cristianos?", disabled=bloquear_inputs)

st.write("8) Propón 3 acciones prácticas que podrías hacer en tu día a día para mejorar la convivencia con personas que tienen creencias diferentes a las tuyas:")
q8_1 = st.text_input("Acción 1:", disabled=bloquear_inputs)
q8_2 = st.text_input("Acción 2:", disabled=bloquear_inputs)
q8_3 = st.text_input("Acción 3:", disabled=bloquear_inputs)

st.markdown("---")
st.subheader("📊 Valoración de la Actividad")
val_funcional = st.slider("1. ¿Qué tan fácil y funcional te pareció usar esta ficha digital?", 1, 5, 5, disabled=bloquear_inputs)
val_interes = st.radio("2. ¿El tema y las actividades te parecieron interesantes?", ["Sí, mucho", "Estuvo bien", "No mucho", "Nada interesante"], horizontal=True, disabled=bloquear_inputs)

st.markdown("---")
st.subheader("📋 Lista de Cotejo (Uso exclusivo del docente)")
st.caption("Estos son los criterios con los que tu profesor evaluará esta ficha:")
st.checkbox("Reconoce el concepto de monoteísmo y la figura de Abraham adecuadamente (Nivel 1).", value=False, disabled=True)
st.checkbox("Identifica de manera asertiva semejanzas entre el judaísmo y el cristianismo (Nivel 2).", value=False, disabled=True)
st.checkbox("Argumenta sólidamente la importancia del respeto a otras creencias desde los valores cristianos (Nivel 3).", value=False, disabled=True)
st.checkbox("Propone acciones prácticas y viables para la convivencia interreligiosa (Nivel 3).", value=False, disabled=True)

st.markdown("---")

if st.button("Generar mi Evidencia en Word"):
    if not nombre.strip() or seccion == "":
        st.error("⚠️ Por favor, ingresa tu nombre y selecciona tu sección para poder identificarte.")
    elif not tiempo_agotado and (
        not q1.strip() or not q3_1.strip() or not q3_2.strip() or not q4.strip() or
        not q5_1.strip() or not q5_2.strip() or not q6.strip() or
        not q7.strip() or not q8_1.strip() or not q8_2.strip() or not q8_3.strip()
    ):
        st.error("⚠️ Aún tienes tiempo. Debes completar las preguntas de **TODOS LOS NIVELES (1, 2 y 3)** antes de descargar tu evidencia.")
    else:
        doc = Document()
        doc.add_heading('FICHA DE RELIGIÓN – SESIÓN 1° AÑO', level=1)
        doc.add_paragraph(f'Estudiante: {nombre} | Sección: {seccion} | Fecha: {fecha}')
        if tiempo_agotado:
            doc.add_paragraph('[Entregado al finalizar el tiempo reglamentario]').bold = True
        doc.add_paragraph('---')
        
        doc.add_heading('NIVEL 1 (FÁCIL)', level=2)
        p1 = doc.add_paragraph()
        p1.add_run('1) Significado de monoteísmo:\n').bold = True
        p1.add_run(f'• {q1}\n\n')
        p1.add_run('2) Importancia de Abraham:\n').bold = True
        p1.add_run(f'• {q2}\n\n')
        p1.add_run('3) Valores aprendidos de Abraham:\n').bold = True
        p1.add_run(f'• {q3_1}\n• {q3_2}\n\n')
        p1.add_run('4) Regla de oro para el diálogo respetuoso:\n').bold = True
        p1.add_run(f'• {q4}')
        
        doc.add_heading('NIVEL 2 (MEDIO)', level=2)
        p2 = doc.add_paragraph()
        p2.add_run('5) Semejanzas entre Judaísmo y Cristianismo:\n').bold = True
        p2.add_run(f'• {q5_1}\n• {q5_2}\n\n')
        p2.add_run('6) Prejuicio a evitar en el salón:\n').bold = True
        p2.add_run(f'• {q6}')
        
        doc.add_heading('NIVEL 3 (DIFÍCIL)', level=2)
        p3 = doc.add_paragraph()
        p3.add_run('7) Reflexión sobre el respeto en los valores cristianos:\n').bold = True
        p3.add_run(f'• {q7}\n\n')
        p3.add_run('8) Acciones prácticas para mejorar la convivencia:\n').bold = True
        p3.add_run(f'• {q8_1}\n• {q8_2}\n• {q8_3}')

        doc.add_heading('Valoración de la Actividad', level=2)
        p4 = doc.add_paragraph()
        p4.add_run('Funcionalidad de la ficha digital: ').bold = True
        p4.add_run(f'{val_funcional} estrellas\n')
        p4.add_run('Interés en el tema: ').bold = True
        p4.add_run(f'{val_interes}')
        
        doc.add_page_break()
        doc.add_heading('Lista de Cotejo - Evaluación del Docente', level=2)
        doc.add_paragraph('[ ] Reconoce el concepto de monoteísmo y la figura de Abraham adecuadamente (Nivel 1).')
        doc.add_paragraph('[ ] Identifica de manera asertiva semejanzas entre el judaísmo y el cristianismo (Nivel 2).')
        doc.add_paragraph('[ ] Argumenta sólidamente la importancia del respeto a otras creencias desde los valores cristianos (Nivel 3).')
        doc.add_paragraph('[ ] Propone acciones prácticas y viables para la convivencia interreligiosa (Nivel 3).')
        doc.add_paragraph('\nNota / Observaciones: ________________________________________________')

        bio = io.BytesIO()
        doc.save(bio)

        if not tiempo_agotado: st.balloons()
        st.success("¡Tu archivo está listo para entregar!")
        st.download_button(label="📥 Descargar Documento Final (.docx)", data=bio.getvalue(), file_name=f"Ficha_Religion_1{seccion}_{nombre.replace(' ', '_')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
