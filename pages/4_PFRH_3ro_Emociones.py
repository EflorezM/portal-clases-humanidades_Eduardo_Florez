import streamlit as st
import time
from docx import Document
import io

st.set_page_config(page_title="Ficha PFRH - Emociones", page_icon="🧠", layout="centered")

# --- SISTEMA DE TIEMPO INTELIGENTE Y BOTÓN GO ---
if 'ficha_iniciada_pfrh' not in st.session_state:
    st.session_state.ficha_iniciada_pfrh = False

if not st.session_state.ficha_iniciada_pfrh:
    st.info("👋 ¡Hola! Tienes 20 minutos para resolver esta ficha. El tiempo comenzará a correr cuando presiones el botón de abajo.")
    col_btn1, col_btn2, col_btn3 = st.columns([1, 2, 1])
    with col_btn2:
        if st.button("🚀 ESTOY LISTO: INICIAR FICHA", use_container_width=True):
            st.session_state.ficha_iniciada_pfrh = True
            st.session_state.inicio_tiempo_pfrh = time.time()
            st.session_state.minutos_asignados_pfrh = 20
            st.rerun()
    st.stop()

segundos_transcurridos = time.time() - st.session_state.inicio_tiempo_pfrh
segundos_restantes = (st.session_state.minutos_asignados_pfrh * 60) - segundos_transcurridos
tiempo_agotado = segundos_restantes <= 0
bloquear_inputs = tiempo_agotado 

# --- MENÚ LATERAL: RELOJ Y TIEMPO EXTRA ---
with st.sidebar:
    st.markdown("### ⏱️ Cronómetro de ficha")
    if not tiempo_agotado:
        minutos = int(segundos_restantes // 60)
        segundos = int(segundos_restantes % 60)
        st.success(f"## {minutos:02d}:{segundos:02d}")
        
        if st.button("➕ Dar 4 min extra"):
            st.session_state.minutos_asignados_pfrh += 4
            st.rerun()
        
        st.caption("Actualiza la página (F5) o interactúa con la ficha para ver el tiempo exacto.")
    else:
        st.error("## 00:00\n⚠️ TIEMPO AGOTADO")
        st.write("Tu ficha ha sido bloqueada. Por favor, descarga tu avance en la parte inferior.")
        
        if st.button("🔓 Desbloquear (Dar 4 min)"):
            st.session_state.minutos_asignados_pfrh += 4
            st.rerun()
# --------------------------------------

st.markdown("### FICHA DE PFRH – SESIÓN 3° AÑO A - B - C")
st.markdown("**Tema:** Las emociones que experimentamos.")
st.markdown("**Prof:** Eduardo Florez Montero / Unidad 1")
st.markdown("---")

col1, col2, col3 = st.columns([3, 1, 2])
with col1: nombre = st.text_input("Estudiante:", disabled=bloquear_inputs)
with col2: seccion = st.selectbox("Sección:", ["", "A", "B", "C"], disabled=bloquear_inputs)
with col3: fecha = st.text_input("Fecha:", disabled=bloquear_inputs)

st.markdown("---")
with st.expander("🎯 META DE HOY", expanded=True): st.write("Reconocer emoción–pensamiento–situación y expresar emociones con respeto.")

st.markdown("---")
st.subheader("📚 Material de Apoyo")
try:
    with open("Presentacion_PFRH.pdf", "rb") as file:
        st.download_button("📄 Descargar Presentación de la Clase (PDF)", data=file, file_name="Presentacion_PFRH.pdf", mime="application/pdf")
except FileNotFoundError:
    st.info("📌 La presentación de la clase se está procesando y aparecerá aquí en breve.")

st.markdown("---")
st.subheader("NIVEL 1 (FÁCIL) – PARA TODOS")
q1 = st.selectbox("1) Haz una pausa y reflexiona: ¿Cuál de estas emociones describe mejor cómo te sientes el día de hoy?", ["", "alegría", "enojo", "tristeza", "miedo", "vergüenza", "calma"], disabled=bloquear_inputs)
q2 = st.text_input("2) Observa tus reacciones físicas. Completa la frase: Cuando siento enojo o frustración, noto que mi cuerpo...", disabled=bloquear_inputs)
q3 = st.text_area("3) Identifica el pensamiento: Escribe un pensamiento frecuente que suele aparecer en tu mente cuando experimentas enojo (recuerda no usar nombres reales de otras personas):", disabled=bloquear_inputs)
q4 = st.text_area("4) Busca una salida sana: Escribe una forma respetuosa y asertiva en la que podrías expresar o canalizar esa emoción sin lastimar a nadie:", disabled=bloquear_inputs)

st.markdown("---")
st.subheader("NIVEL 2 (MEDIO) – RETO 1")
st.write("5) Analiza el siguiente caso: Le escribes un mensaje importante a un amigo y te deja 'en visto' sin responder. Completa los espacios identificando qué emoción sentirías, qué pensarías y cuál sería tu respuesta:")
col_c1, col_c2, col_c3 = st.columns(3)
with col_c1: q5_emo = st.text_input("¿Qué emoción sentirías?:", disabled=bloquear_inputs)
with col_c2: q5_pen = st.text_input("¿Qué pensamiento cruzaría tu mente?:", disabled=bloquear_inputs)
with col_c3: q5_res = st.text_input("¿Cuál sería una respuesta respetuosa?:", disabled=bloquear_inputs)

st.markdown("---")
st.subheader("NIVEL 3 (DIFÍCIL) – RETO 2")
q6 = st.text_area("6) Reflexiona profundo (5 líneas): ¿Cuáles crees que son las consecuencias emocionales y sociales si te guardas todo lo que sientes y terminas 'explotando' impulsivamente en tus redes sociales?", disabled=bloquear_inputs)
st.write("7) Crea dos frases cortas y positivas de 'autocuidado' emocional que podrías repetirte a ti mismo/a en momentos de mucho estrés o tristeza:")
q7_1 = st.text_input("Frase de autocuidado 1:", disabled=bloquear_inputs)
q7_2 = st.text_input("Frase de autocuidado 2:", disabled=bloquear_inputs)

st.markdown("---")
# --- VALORACIÓN DEL ESTUDIANTE ---
st.subheader("📊 Valoración de la Actividad")
val_funcional = st.slider("1. ¿Qué tan fácil y funcional te pareció usar esta ficha digital?", 1, 5, 5, disabled=bloquear_inputs)
val_interes = st.radio("2. ¿El tema y las actividades te parecieron interesantes?", ["Sí, mucho", "Estuvo bien", "No mucho", "Nada interesante"], horizontal=True, disabled=bloquear_inputs)

st.markdown("---")
# --- LISTA DE COTEJO PARA EL DOCENTE ---
st.subheader("📋 Lista de Cotejo (Uso exclusivo del docente)")
st.caption("Estos son los criterios con los que tu profesor evaluará esta ficha:")
st.checkbox("Identifica correctamente sus emociones y reacciones físicas (Nivel 1).", value=False, disabled=True)
st.checkbox("Relaciona asertivamente emoción, pensamiento y respuesta en un caso cotidiano (Nivel 2).", value=False, disabled=True)
st.checkbox("Reflexiona con profundidad sobre las consecuencias de no gestionar emociones (Nivel 3).", value=False, disabled=True)
st.checkbox("Propone frases de autocuidado emocional coherentes y positivas (Nivel 3).", value=False, disabled=True)

st.markdown("---")

if st.button("Generar mi Evidencia en Word"):
    if not nombre.strip() or seccion == "":
        st.error("⚠️ Por favor, ingresa tu nombre y selecciona tu sección para poder identificarte.")
        
    elif not tiempo_agotado and (
        q1 == "" or not q2.strip() or not q3.strip() or not q4.strip() or
        not q5_emo.strip() or not q5_pen.strip() or not q5_res.strip() or
        not q6.strip() or not q7_1.strip() or not q7_2.strip()
    ):
        st.error("⚠️ Aún tienes tiempo. Debes completar las preguntas de **TODOS LOS NIVELES (1, 2 y 3)** antes de descargar tu evidencia.")
        
    else:
        doc = Document()
        
        doc.add_heading('FICHA DE PFRH – SESIÓN 3° AÑO', level=1)
        doc.add_paragraph(f'Estudiante: {nombre} | Sección: {seccion} | Fecha: {fecha}')
        if tiempo_agotado:
            doc.add_paragraph('[Entregado al finalizar el tiempo reglamentario]').bold = True
        doc.add_paragraph('---')
        
        doc.add_heading('NIVEL 1 (FÁCIL)', level=2)
        p1 = doc.add_paragraph()
        p1.add_run('1) Emoción del día:\n').bold = True
        p1.add_run(f'• {q1}\n\n')
        p1.add_run('2) Reacción del cuerpo al enojo:\n').bold = True
        p1.add_run(f'• {q2}\n\n')
        p1.add_run('3) Pensamiento frecuente:\n').bold = True
        p1.add_run(f'• {q3}\n\n')
        p1.add_run('4) Forma respetuosa de canalizar la emoción:\n').bold = True
        p1.add_run(f'• {q4}')
        
        doc.add_heading('NIVEL 2 (MEDIO)', level=2)
        p2 = doc.add_paragraph()
        p2.add_run('5) Caso "Me dejan en visto":\n').bold = True
        p2.add_run(f'• Emoción: {q5_emo}\n')
        p2.add_run(f'• Pensamiento: {q5_pen}\n')
        p2.add_run(f'• Respuesta: {q5_res}')
        
        doc.add_heading('NIVEL 3 (DIFÍCIL)', level=2)
        p3 = doc.add_paragraph()
        p3.add_run('6) Consecuencias de explotar en redes:\n').bold = True
        p3.add_run(f'• {q6}\n\n')
        p3.add_run('7) Frases de autocuidado:\n').bold = True
        p3.add_run(f'• {q7_1}\n• {q7_2}')
        
        doc.add_heading('Valoración de la Actividad', level=2)
        p4 = doc.add_paragraph()
        p4.add_run('Funcionalidad de la ficha digital: ').bold = True
        p4.add_run(f'{val_funcional} estrellas\n')
        p4.add_run('Interés en el tema: ').bold = True
        p4.add_run(f'{val_interes}')
        
        doc.add_page_break()
        
        doc.add_heading('Lista de Cotejo - Evaluación del Docente', level=2)
        doc.add_paragraph('[ ] Identifica correctamente sus emociones y reacciones físicas (Nivel 1).')
        doc.add_paragraph('[ ] Relaciona asertivamente emoción, pensamiento y respuesta en un caso cotidiano (Nivel 2).')
        doc.add_paragraph('[ ] Reflexiona con profundidad sobre las consecuencias de no gestionar emociones (Nivel 3).')
        doc.add_paragraph('[ ] Propone frases de autocuidado emocional coherentes y positivas (Nivel 3).')
        doc.add_paragraph('\nNota / Observaciones: ________________________________________________')
        
        bio = io.BytesIO()
        doc.save(bio)
        
        if not tiempo_agotado: st.balloons()
            
        st.success("¡Tu archivo está listo para entregar!")
        st.download_button(
            label="📥 Descargar Documento Final (.docx)", 
            data=bio.getvalue(), 
            file_name=f"Ficha_PFRH_3{seccion}_{nombre.replace(' ', '_')}.docx", 
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
