import streamlit as st
import time
from docx import Document
import io
import streamlit.components.v1 as components

st.set_page_config(page_title="Plan Lector - Chepeconde", page_icon="📖", layout="centered")

# --- 1. LA BÓVEDA DE SEGURIDAD ---
if 'autenticado_chepe' not in st.session_state:
    st.session_state.autenticado_chepe = False

if not st.session_state.autenticado_chepe:
    st.title("📖 Plan Lector Transmedia")
    st.subheader("Obra: El Pueblo de Chepeconde")
    st.write("Autor: Eduardo Florez Montero")
    st.markdown("---")
    st.info("🔒 **Acceso Exclusivo:** Escribe la Clave de Colegio que tu profesor te indicó o que vino en tu libro.")
    
    clave = st.text_input("🔑 Clave de Acceso:", type="password")
    
    if st.button("Desbloquear Evaluación", use_container_width=True):
        if clave == "CHEPECONDE-2026": 
            st.session_state.autenticado_chepe = True
            st.rerun()
        else:
            st.error("❌ Clave incorrecta. Verifica e intenta nuevamente.")
    st.stop()

# --- 2. SISTEMA DE TIEMPO INTELIGENTE ---
if 'ficha_iniciada_chepe' not in st.session_state:
    st.session_state.ficha_iniciada_chepe = False

if not st.session_state.ficha_iniciada_chepe:
    st.success("✅ ¡Clave aceptada! Bienvenido al ecosistema de El Pueblo de Chepeconde.")
    
    # --- ESPACIO PARA LA PORTADA DEL LIBRO ---
    try:
        st.image("portada_chepeconde.jpg", caption="El Pueblo de Chepeconde", width=250)
    except FileNotFoundError:
        st.caption("(Aquí aparecerá la portada de tu libro cuando subas la imagen 'portada_chepeconde.jpg' a GitHub)")
        
    st.info("👋 Tienes 20 minutos para resolver tu evaluación de comprensión. Presiona el botón cuando estés listo.")
    col_btn1, col_btn2, col_btn3 = st.columns([1, 2, 1])
    with col_btn2:
        if st.button("🚀 INICIAR EVALUACIÓN", use_container_width=True):
            st.session_state.ficha_iniciada_chepe = True
            st.session_state.inicio_tiempo_chepe = time.time()
            st.session_state.minutos_asignados_chepe = 20
            st.rerun()
    st.stop()

segundos_transcurridos = time.time() - st.session_state.inicio_tiempo_chepe
segundos_restantes = (st.session_state.minutos_asignados_chepe * 60) - segundos_transcurridos
tiempo_agotado = segundos_restantes <= 0
bloquear_inputs = tiempo_agotado 

# --- MENÚ LATERAL: RELOJ VISUAL ---
with st.sidebar:
    st.markdown("### ⏱️ Tiempo Restante")
    if not tiempo_agotado:
        reloj_html = f"""
        <div style="background-color: #f0f2f6; padding: 15px; border-radius: 10px; text-align: center; border: 2px solid #e0e4eb;">
            <h2 id="reloj" style="margin: 0; color: #2ecc71; font-family: monospace; font-size: 38px;">--:--</h2>
        </div>
        <script>
            var tiempo = {int(segundos_restantes)};
            var display = document.getElementById('reloj');
            var intervalo = setInterval(function() {{
                if (tiempo <= 0) {{
                    clearInterval(intervalo);
                    display.innerHTML = "00:00";
                    display.style.color = "#e74c3c";
                }} else {{
                    var min = Math.floor(tiempo / 60).toString().padStart(2, '0');
                    var sec = (tiempo % 60).toString().padStart(2, '0');
                    display.innerHTML = min + ":" + sec;
                    if (tiempo <= 300) display.style.color = "#f39c12"; 
                    if (tiempo <= 60) display.style.color = "#e74c3c"; 
                    tiempo--;
                }}
            }}, 1000);
        </script>
        """
        components.html(reloj_html, height=85)
        if st.button("➕ Dar 4 min extra", use_container_width=True):
            st.session_state.minutos_asignados_chepe += 4
            st.rerun()
    else:
        st.error("## 00:00\n⚠️ TIEMPO AGOTADO")
        st.write("Tu ficha ha sido bloqueada. Descarga tu avance abajo.")
        if st.button("🔓 Desbloquear (Dar 4 min extra)", use_container_width=True):
            st.session_state.minutos_asignados_chepe += 4
            st.rerun()

# --- 3. CONTENIDO LITERARIO ---
st.markdown("### 📖 EL PUEBLO DE CHEPECONDE")
st.markdown("**Autor:** Eduardo Florez Montero | **Categoría:** Pre-Adolescentes")
st.markdown("---")

col1, col2, col3 = st.columns([3, 1, 2])
with col1: nombre = st.text_input("Estudiante:", disabled=bloquear_inputs)
with col2: seccion = st.selectbox("Sección:", ["", "A", "B", "C"], disabled=bloquear_inputs)
with col3: fecha = st.text_input("Fecha:", disabled=bloquear_inputs)

st.markdown("---")
# --- INTEGRACIÓN DEL VIDEO (REEL/BOOKTRAILER) ---
st.subheader("🎬 Antes de empezar: Resumen Visual")
st.video("https://www.youtube.com/watch?v=123456789") # <-- PON AQUÍ EL LINK DE TU VIDEO CUANDO LO TENGAS
st.caption("Responde basándote en la lectura de tu libro.")
st.markdown("---")

st.subheader("NIVEL 1: COMPRENSIÓN LITERAL")
q1 = st.text_input("1) ¿Quién es el personaje principal y cuál es su mayor característica?", disabled=bloquear_inputs)
q2 = st.text_area("2) Según el inicio del libro, describe brevemente cómo es El Pueblo de Chepeconde:", disabled=bloquear_inputs)

st.markdown("---")
st.subheader("NIVEL 2: ANÁLISIS E INFERENCIA")
q3 = st.text_area("3) ¿Por qué crees que el protagonista tomó esa decisión difícil? ¿Qué hubieras hecho tú?", disabled=bloquear_inputs)
q4 = st.text_input("4) Identifica el conflicto principal de la historia. ¿Era un problema interno o externo?", disabled=bloquear_inputs)

st.markdown("---")
st.subheader("NIVEL 3: VALORACIÓN CRÍTICA")
q5 = st.text_area("5) Reflexión profunda (5 líneas): ¿Qué enseñanza o valor nos deja la historia de Chepeconde para la sociedad actual?", disabled=bloquear_inputs)

st.markdown("---")
st.subheader("📋 Lista de Cotejo (Uso exclusivo del docente)")
st.checkbox("Identifica correctamente personajes y escenarios (Nivel 1).", value=False, disabled=True)
st.checkbox("Analiza los motivos y conflictos internos de los personajes (Nivel 2).", value=False, disabled=True)
st.checkbox("Argumenta sólidamente una reflexión ética basada en la obra (Nivel 3).", value=False, disabled=True)

st.markdown("---")

if st.button("Generar mi Evidencia en Word"):
    if not nombre.strip() or seccion == "":
        st.error("⚠️ Por favor, ingresa tu nombre y selecciona tu sección.")
    elif not tiempo_agotado and (not q1.strip() or not q2.strip() or not q3.strip() or not q4.strip() or not q5.strip()):
        st.error("⚠️ Aún tienes tiempo. Debes completar TODAS las preguntas.")
    else:
        doc = Document()
        doc.add_heading('PLAN LECTOR – EL PUEBLO DE CHEPECONDE', level=1)
        doc.add_paragraph(f'Estudiante: {nombre} | Sección: {seccion} | Fecha: {fecha}')
        if tiempo_agotado:
            doc.add_paragraph('[Entregado al finalizar el tiempo reglamentario]').bold = True
        doc.add_paragraph('---')
        
        doc.add_heading('NIVEL 1 (Literal)', level=2)
        p1 = doc.add_paragraph()
        p1.add_run('1) Personaje principal:\n').bold = True
        p1.add_run(f'• {q1}\n\n')
        p1.add_run('2) Descripción del pueblo:\n').bold = True
        p1.add_run(f'• {q2}')
        
        doc.add_heading('NIVEL 2 (Inferencia)', level=2)
        p2 = doc.add_paragraph()
        p2.add_run('3) Decisión del personaje:\n').bold = True
        p2.add_run(f'• {q3}\n\n')
        p2.add_run('4) Conflicto principal:\n').bold = True
        p2.add_run(f'• {q4}')
        
        doc.add_heading('NIVEL 3 (Crítico)', level=2)
        p3 = doc.add_paragraph()
        p3.add_run('5) Reflexión final:\n').bold = True
        p3.add_run(f'• {q5}')
        
        doc.add_page_break()
        doc.add_heading('Lista de Cotejo - Evaluación del Docente', level=2)
        doc.add_paragraph('[ ] Identifica correctamente personajes y escenarios (Nivel 1).')
        doc.add_paragraph('[ ] Analiza los motivos y conflictos internos de los personajes (Nivel 2).')
        doc.add_paragraph('[ ] Argumenta sólidamente una reflexión ética basada en la obra (Nivel 3).')
        doc.add_paragraph('\nNota / Observaciones: ________________________________________________')

        bio = io.BytesIO()
        doc.save(bio)
        
        if not tiempo_agotado: st.balloons()
        st.success("¡Tu archivo está listo para entregar!")
        st.download_button(label="📥 Descargar Documento Final (.docx)", data=bio.getvalue(), file_name=f"Chepeconde_{seccion}_{nombre.replace(' ', '_')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
