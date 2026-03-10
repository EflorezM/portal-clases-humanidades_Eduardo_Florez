import streamlit as st
import time
from docx import Document
import io
import streamlit.components.v1 as components

st.set_page_config(page_title="Ficha Cívica 3ro - Dimensiones", page_icon="⚖️", layout="centered")

# --- SISTEMA DE TIEMPO INTELIGENTE Y BOTÓN GO ---
if 'ficha_iniciada_civica3' not in st.session_state:
    st.session_state.ficha_iniciada_civica3 = False

if not st.session_state.ficha_iniciada_civica3:
    st.info("👋 ¡Hola! Tienes 20 minutos para resolver esta ficha. El tiempo comenzará a correr cuando presiones el botón de abajo.")
    col_btn1, col_btn2, col_btn3 = st.columns([1, 2, 1])
    with col_btn2:
        if st.button("🚀 ESTOY LISTO: INICIAR FICHA", use_container_width=True):
            st.session_state.ficha_iniciada_civica3 = True
            st.session_state.inicio_tiempo_civica3 = time.time()
            st.session_state.minutos_asignados_civica3 = 20
            st.rerun()
    st.stop()

segundos_transcurridos = time.time() - st.session_state.inicio_tiempo_civica3
segundos_restantes = (st.session_state.minutos_asignados_civica3 * 60) - segundos_transcurridos
tiempo_agotado = segundos_restantes <= 0
bloquear_inputs = tiempo_agotado 

# --- MENÚ LATERAL: RELOJ VISUAL EN TIEMPO REAL ---
with st.sidebar:
    st.markdown("### ⏱️ Cronómetro de ficha")
    if not tiempo_agotado:
        # Reloj visual dinámico con JavaScript
        reloj_html = f"""
        <div style="background-color: #f0f2f6; padding: 15px; border-radius: 10px; text-align: center; border: 2px solid #e0e4eb;">
            <h2 id="reloj" style="margin: 0; color: #2ecc71; font-family: monospace; font-size: 38px;">--:--</h2>
        </div>
        <script>
            var tiempo = {int(segundos_restantes)};
            var display = document.getElementById('reloj');
            var intervalo = setInterval(function() {{
                if (tiempo <= 0) {{
                    clearInterval(intervalo);
                    display.innerHTML = "00:00";
                    display.style.color = "#e74c3c";
                }} else {{
                    var min = Math.floor(tiempo / 60).toString().padStart(2, '0');
                    var sec = (tiempo % 60).toString().padStart(2, '0');
                    display.innerHTML = min + ":" + sec;
                    if (tiempo <= 300) display.style.color = "#f39c12"; // Amarillo últimos 5 min
                    if (tiempo <= 60) display.style.color = "#e74c3c"; // Rojo último minuto
                    tiempo--;
                }}
            }}, 1000);
        </script>
        """
        components.html(reloj_html, height=85)
        
        if st.button("➕ Dar 4 min extra", use_container_width=True):
            st.session_state.minutos_asignados_civica3 += 4
            st.rerun()
        st.caption("El reloj avanza en tiempo real. Al llegar a cero, haz clic en cualquier parte para bloquear y descargar tu avance.")
    else:
        st.error("## 00:00\n⚠️ TIEMPO AGOTADO")
        st.write("Tu ficha ha sido bloqueada. Por favor, descarga tu avance en la parte inferior.")
        
        if st.button("🔓 Desbloquear (Dar 4 min extra)", use_container_width=True):
            st.session_state.minutos_asignados_civica3 += 4
            st.rerun()
# --------------------------------------

st.markdown("### FICHA DE CÍVICA – SESIÓN 3° AÑO A - B - C")
st.markdown("**Tema:** Las dimensiones de la realidad humana.")
st.markdown("**Prof:** Eduardo Florez Montero / Unidad 1")
st.markdown("---")

col1, col2, col3 = st.columns([3, 1, 2])
with col1: nombre = st.text_input("Estudiante:", disabled=bloquear_inputs)
with col2: seccion = st.selectbox("Sección:", ["", "A", "B", "C"], disabled=bloquear_inputs)
with col3: fecha = st.text_input("Fecha:", disabled=bloquear_inputs)

st.markdown("---")
with st.expander("🎯 META DE HOY", expanded=True): st.write("Reconocer dimensiones humanas (biológica, social, cultural y estética) y relacionarlas con identidad y convivencia digital.")

st.markdown("---")
st.subheader("📚 Material de Apoyo")
try:
    with open("Presentacion_Civica_3ro.pdf", "rb") as file:
        st.download_button("📄 Descargar Presentación de la Clase (PDF)", data=file, file_name="Presentacion_Civica_3ro.pdf", mime="application/pdf")
except FileNotFoundError:
    st.info("📌 La presentación de la clase se está procesando y aparecerá aquí en breve.")

st.markdown("---")
st.subheader("NIVEL 1 (FÁCIL) – PARA TODOS")
st.write("¡Empecemos! Vamos a reconocer de qué estamos hechos y cómo interactuamos.")
q1 = st.text_input("1) Según lo que hemos conversado, ¿cómo explicarías con tus propias palabras qué es la dimensión social?", disabled=bloquear_inputs)

st.write("2) Escribe un ejemplo breve de tu vida diaria para cada una de estas dimensiones:")
col_d1, col_d2 = st.columns(2)
with col_d1:
    q2_bio = st.text_input("Biológica:", disabled=bloquear_inputs)
    q2_cul = st.text_input("Cultural:", disabled=bloquear_inputs)
with col_d2:
    q2_soc = st.text_input("Social:", disabled=bloquear_inputs)
    q2_est = st.text_input("Estética:", disabled=bloquear_inputs)

q3 = st.radio("3) Si alguien participa en un 'reto' de internet que pone en riesgo su integridad física, ¿qué dimensión se ve más afectada directamente?", ["Biológica", "Cultural", "Estética"], horizontal=True, disabled=bloquear_inputs)
q4 = st.text_input("4) Escribe una decisión responsable que puedes tomar hoy mismo para cuidar tu bienestar en redes sociales:", disabled=bloquear_inputs)

st.markdown("---")
st.subheader("NIVEL 2 (MEDIO) – RETO 1")
st.write("5) Analiza el siguiente caso: Un reto peligroso se vuelve viral en internet y muchos jóvenes intentan copiarlo. ¿Cómo crees que este reto impactaría negativamente en cada una de las dimensiones de un estudiante? Escribe una idea para cada una:")
col_c1, col_c2 = st.columns(2)
with col_c1:
    q5_bio = st.text_input("Impacto en su dimensión Biológica:", disabled=bloquear_inputs)
    q5_soc = st.text_input("Impacto en su dimensión Social:", disabled=bloquear_inputs)
with col_c2:
    q5_cul = st.text_input("Impacto en su dimensión Cultural:", disabled=bloquear_inputs)
    q5_est = st.text_input("Impacto en su dimensión Estética:", disabled=bloquear_inputs)

st.markdown("---")
st.subheader("NIVEL 3 (DIFÍCIL) – RETO 2")
q6 = st.text_area("6) Reflexión profunda (aprox. 5 líneas): De las cuatro dimensiones estudiadas, ¿cuál crees que influye más en la forma en que construyes tu propia identidad en esta etapa de tu vida? ¿Por qué?", disabled=bloquear_inputs)

st.write("7) Propón 3 criterios o 'filtros' personales que deberías usar antes de elegir qué contenidos consumir o compartir en tus redes:")
q7_1 = st.text_input("Criterio 1:", disabled=bloquear_inputs)
q7_2 = st.text_input("Criterio 2:", disabled=bloquear_inputs)
q7_3 = st.text_input("Criterio 3:", disabled=bloquear_inputs)

st.markdown("---")
# --- VALORACIÓN DEL ESTUDIANTE ---
st.subheader("📊 Valoración de la Actividad")
val_funcional = st.slider("1. ¿Qué tan fácil y funcional te pareció usar esta ficha digital?", 1, 5, 5, disabled=bloquear_inputs)
val_interes = st.radio("2. ¿El tema y las actividades te parecieron interesantes?", ["Sí, mucho", "Estuvo bien", "No mucho", "Nada interesante"], horizontal=True, disabled=bloquear_inputs)

st.markdown("---")
# --- LISTA DE COTEJO PARA EL DOCENTE ---
st.subheader("📋 Lista de Cotejo (Uso exclusivo del docente)")
st.caption("Estos son los criterios con los que tu profesor evaluará esta ficha:")
st.checkbox("Define y ejemplifica correctamente las cuatro dimensiones de la realidad humana (Nivel 1).", value=False, disabled=True)
st.checkbox("Identifica decisiones responsables para el uso de redes sociales (Nivel 1).", value=False, disabled=True)
st.checkbox("Analiza integralmente cómo las situaciones de riesgo afectan todas las dimensiones de la persona (Nivel 2).", value=False, disabled=True)
st.checkbox("Argumenta críticamente la influencia de las dimensiones en su identidad y establece filtros de contenido (Nivel 3).", value=False, disabled=True)

st.markdown("---")

if st.button("Generar mi Evidencia en Word"):
    if not nombre.strip() or seccion == "":
        st.error("⚠️ Por favor, ingresa tu nombre y selecciona tu sección para poder identificarte.")
        
    elif not tiempo_agotado and (
        not q1.strip() or not q2_bio.strip() or not q2_soc.strip() or not q2_cul.strip() or not q2_est.strip() or not q4.strip() or
        not q5_bio.strip() or not q5_soc.strip() or not q5_cul.strip() or not q5_est.strip() or
        not q6.strip() or not q7_1.strip() or not q7_2.strip() or not q7_3.strip()
    ):
        st.error("⚠️ Aún tienes tiempo. Debes completar las preguntas de **TODOS LOS NIVELES (1, 2 y 3)** antes de descargar tu evidencia.")
        
    else:
        doc = Document()
        
        doc.add_heading('FICHA DE CÍVICA – SESIÓN 3° AÑO', level=1)
        doc.add_paragraph(f'Estudiante: {nombre} | Sección: {seccion} | Fecha: {fecha}')
        if tiempo_agotado:
            doc.add_paragraph('[Entregado al finalizar el tiempo reglamentario]').bold = True
        doc.add_paragraph('---')
        
        doc.add_heading('NIVEL 1 (FÁCIL)', level=2)
        p1 = doc.add_paragraph()
        p1.add_run('1) Definición de dimensión social:\n').bold = True
        p1.add_run(f'• {q1}\n\n')
        p1.add_run('2) Ejemplos por dimensión:\n').bold = True
        p1.add_run(f'• Biológica: {q2_bio}\n• Social: {q2_soc}\n• Cultural: {q2_cul}\n• Estética: {q2_est}\n\n')
        p1.add_run('3) Dimensión afectada por reto peligroso:\n').bold = True
        p1.add_run(f'• {q3}\n\n')
        p1.add_run('4) Decisión responsable en redes:\n').bold = True
        p1.add_run(f'• {q4}')
        
        doc.add_heading('NIVEL 2 (MEDIO)', level=2)
        p2 = doc.add_paragraph()
        p2.add_run('5) Impacto de un reto viral por dimensión:\n').bold = True
        p2.add_run(f'• Biológico: {q5_bio}\n')
        p2.add_run(f'• Social: {q5_soc}\n')
        p2.add_run(f'• Cultural: {q5_cul}\n')
        p2.add_run(f'• Estético: {q5_est}')
        
        doc.add_heading('NIVEL 3 (DIFÍCIL)', level=2)
        p3 = doc.add_paragraph()
        p3.add_run('6) Dimensión de mayor influencia en la identidad:\n').bold = True
        p3.add_run(f'• {q6}\n\n')
        p3.add_run('7) Criterios para elegir contenidos:\n').bold = True
        p3.add_run(f'• {q7_1}\n• {q7_2}\n• {q7_3}')
        
        doc.add_heading('Valoración de la Actividad', level=2)
        p4 = doc.add_paragraph()
        p4.add_run('Funcionalidad de la ficha digital: ').bold = True
        p4.add_run(f'{val_funcional} estrellas\n')
        p4.add_run('Interés en el tema: ').bold = True
        p4.add_run(f'{val_interes}')
        
        doc.add_page_break()
        
        doc.add_heading('Lista de Cotejo - Evaluación del Docente', level=2)
        doc.add_paragraph('[ ] Define y ejemplifica correctamente las cuatro dimensiones de la realidad humana (Nivel 1).')
        doc.add_paragraph('[ ] Identifica decisiones responsables para el uso de redes sociales (Nivel 1).')
        doc.add_paragraph('[ ] Analiza integralmente cómo las situaciones de riesgo afectan todas las dimensiones de la persona (Nivel 2).')
        doc.add_paragraph('[ ] Argumenta críticamente la influencia de las dimensiones en su identidad y establece filtros de contenido (Nivel 3).')
        doc.add_paragraph('\nNota / Observaciones: ________________________________________________')
        
        bio = io.BytesIO()
        doc.save(bio)
        
        if not tiempo_agotado: st.balloons()
            
        st.success("¡Tu archivo está listo para entregar!")
        st.download_button(
            label="📥 Descargar Documento Final (.docx)", 
            data=bio.getvalue(), 
            file_name=f"Ficha_Civica_3{seccion}_{nombre.replace(' ', '_')}.docx", 
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
